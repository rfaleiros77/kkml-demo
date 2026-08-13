#!/usr/bin/env python3
"""Generate docs/OPERATIONS_RUNBOOK.docx — the project's living runbook.

Single source of truth for the document: edit THIS file, rerun, commit.
The .docx is generated output — hand edits are lost by design.

Structure contract (see the runbook skill):
  title page → 1. Purpose → 2. Environments → procedure sections (prose)
  → validation log → Change Log → FINAL section: Quick Reference step tables.
Every procedure change updates: its prose, its step table, one changelog row.
"""
import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------------------------------------------------- CONFIG (edit me)
PROJECT = "KKML0 Costing Reports — Interactive Demo"
SUBTITLE = "Operations Runbook — living operations manual"
VERSION = "1.0"
DATE = "2026-08-13"
OUT = "docs/OPERATIONS_RUNBOOK.docx"
# ---------------------------------------------------------------------------

FONT = "Arial"
ACC = RGBColor(0x0A, 0x5C, 0x99)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
CODE_RED = RGBColor(0x8B, 0x00, 0x00)

doc = Document()
st = doc.styles["Normal"]
st.font.name = FONT
st.font.size = Pt(10.5)
st.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
for lvl, size, color in (("Heading 1", 15, ACC), ("Heading 2", 12, DARK)):
    s = doc.styles[lvl]
    s.font.name = FONT; s.font.size = Pt(size)
    s.font.bold = True; s.font.color.rgb = color


def rich(par, text):
    """Inline markup: **bold** and `code`."""
    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", str(text)):
        if not part:
            continue
        if part.startswith("**"):
            r = par.add_run(part[2:-2]); r.bold = True; r.font.name = FONT
        elif part.startswith("`"):
            r = par.add_run(part[1:-1])
            r.font.name = "Courier New"; r.font.size = Pt(9.5)
            r.font.color.rgb = CODE_RED
        else:
            r = par.add_run(part); r.font.name = FONT
    return par


def p(text, space=6):
    par = doc.add_paragraph()
    par.paragraph_format.space_after = Pt(space)
    return rich(par, text)


def bullet(text):
    par = doc.add_paragraph(style="List Bullet")
    par.paragraph_format.space_after = Pt(3)
    return rich(par, text)


def step(n, text):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(0.6)
    par.paragraph_format.space_after = Pt(4)
    r = par.add_run(f"{n}.  "); r.bold = True; r.font.name = FONT
    return rich(par, text)


def code(text):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(0.6)
    par.paragraph_format.space_after = Pt(6)
    r = par.add_run(text)
    r.font.name = "Courier New"; r.font.size = Pt(9)
    pr = par._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "F2F2F0")
    pr.append(shd)
    return par


def note(text):
    par = doc.add_paragraph()
    par.paragraph_format.left_indent = Cm(0.4)
    par.paragraph_format.space_after = Pt(6)
    r = par.add_run("NOTE: "); r.bold = True
    r.font.color.rgb = ACC; r.font.name = FONT
    rich(par, text)
    pr = par._p.get_or_add_pPr()
    bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single"); left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "4"); left.set(qn("w:color"), "0A5C99")
    bdr.append(left); pr.append(bdr)
    return par


def table(headers, rows, widths_cm):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.width = Cm(widths_cm[j])
        r = cell.paragraphs[0].add_run(h)
        r.bold = True; r.font.name = FONT; r.font.size = Pt(9.5)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "E8F1F8")
        cell._tc.get_or_add_tcPr().append(shd)
    for i, row in enumerate(rows, start=1):
        for j, c in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.width = Cm(widths_cm[j])
            rich(cell.paragraphs[0], c)
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def h1(text):
    doc.add_heading(text, level=1)


def h2(text):
    doc.add_heading(text, level=2)


# --------------------------------------------------------------- title page
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(220)
r = tp.add_run(f"{PROJECT} — Operations Runbook")
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = ACC; r.font.name = FONT
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run(SUBTITLE); r.font.size = Pt(13); r.font.name = FONT
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = meta.add_run(f"Version {VERSION} · {DATE} · maintained in the project repository ({OUT})")
r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66); r.font.name = FONT
doc.add_page_break()

# ----------------------------------------------------------------- sections
h1("1. Purpose and maintenance")
p("This runbook records every operational procedure of the KKML0 Costing "
  "Reports interactive demo so any of them can be repeated — on any machine, "
  "by any reader — without rediscovering the steps. The demo is a public, "
  "static showcase of the KKML0 cost component matrix reports (SAP Material "
  "Ledger), served from GitHub Pages. Every figure and label is fictitious "
  "sample data.")
bullet("**Living document**: updated whenever a procedure changes or a new one is learned. Each update adds a Change Log row and bumps the version.")
bullet("**Structure rule**: prose sections tell the story (context, findings, guidance); the FINAL section — Quick Reference Step Tables — holds the executable steps. Every procedure change must also refresh its step table.")
bullet("Generated by `tools/build_runbook.py` — never edit the .docx by hand.")
note("This folder (`_arquivo/kkml-demo-backup`) is a **backup snapshot** of the published site. It contains only built artifacts (23 files, no source). The live repository is `rfaleiros77/kkml-demo`; both share the same `origin` remote. The frontend source that produces the bundles is NOT in this repo.")

h1("2. Environments")
table(["Item", "Value"], [
    ["Git repository", "`github.com/rfaleiros77/kkml-demo` (private)"],
    ["Live site", "`https://rfaleiros77.github.io/kkml-demo/`"],
    ["Hosting", "GitHub Pages — legacy build, branch `main`, path `/` (root)"],
    ["Publish trigger", "Any push to `main` → Pages rebuilds automatically"],
    ["Local machine", "Mac mini / MacBook (macOS), Python 3 for tooling"],
    ["Credentials", "None stored. GitHub push uses the machine's `gh` auth."],
], [4.5, 12.0])
p("The demo is fully static — there is no backend, database, or live OData "
  "call at runtime. The Fiori manifests reference OData service paths "
  "(`/sap/opu/odata4/...`) only as design metadata; the running demo shows "
  "synthetic data baked into the bundles.")

h1("3. Site structure — the four flavors")
p("The landing page (`index.html`) links to four self-contained sub-apps, in "
  "two frontend styles, each with an \"actual\" and a \"released\" variant of "
  "the cost matrix:")
table(["Folder", "Style", "Build", "Notes"], [
    ["`costmatrix/`", "Custom UI", "Vite (hashed assets)", "Actual cost matrix"],
    ["`releasedmatrix/`", "Custom UI", "Vite (hashed assets)", "Released cost matrix"],
    ["`fiori-actual/`", "SAP Fiori / UI5", "UI5 (Component-preload)", "Actual, Fiori look"],
    ["`fiori-released/`", "SAP Fiori / UI5", "UI5 (Component-preload)", "Released, Fiori look"],
], [3.6, 3.4, 4.5, 5.0])
note("Asset filenames in the custom apps are content-hashed (e.g. "
     "`index-BaZiN-K7.js`). `costmatrix` and `releasedmatrix` share the same "
     "CSS bundle (`index-BmCR6MLI.css`) but have different JS bundles.")

h1("4. Source protection — the .gitignore rule")
p("This is a **public demo**. Debug bundles and source maps must never be "
  "published, because they reconstruct the original, readable source code for "
  "anyone who downloads them. On 12/08/2026 the `.gitignore` was set to block "
  "them permanently:")
code("*.map\n*-dbg.js\n*-dbg.controller.js\n*-dbg.view.xml")
bullet("`*.map` — source maps map minified bundles back to original source.")
bullet("`*-dbg.js` / `*-dbg.controller.js` / `*-dbg.view.xml` — UI5 ships un-minified \"-dbg\" twins of every file; together they are the full source.")
p("**Before every deploy**, confirm none of these slipped in. As of "
  "13/08/2026 the working tree and the published commit contain zero matches "
  "(verified — see validation log).")

h1("5. Local preview")
p("To view the demo before pushing, serve the folder over HTTP from its root "
  "(opening `index.html` via `file://` breaks the sub-app relative links). "
  "Any static server works; Python's built-in one needs no install:")
code("cd ~/dev/_arquivo/kkml-demo-backup\npython3 -m http.server 8080\n# then open http://localhost:8080/")
p("Click through all four cards from the landing page and confirm each "
  "sub-app renders its table with data. Stop the server with Ctrl+C.")

h1("6. Deploy (publish to GitHub Pages)")
p("Publishing is a plain `git push` to `main` of the `kkml-demo` repo — "
  "GitHub Pages (legacy build) redeploys the root of `main` within a minute or "
  "two. There is no build step on the server and no separate `gh-pages` branch.")
step(1, "Confirm no debug/map files are staged (section 4).")
step(2, "Commit the change with a clear message.")
step(3, "`git push origin main`.")
step(4, "Wait ~1–2 min, then hard-refresh the live URL and re-test the four flavors.")
note("This backup folder pushes to the same `origin` as the live site. Treat a "
     "push from here as a real publish — do not push experimental edits from "
     "the backup unless you intend them to go live.")

h1("7. Validation log")
table(["Date", "Test", "Result"], [
    ["2026-08-13", "No `*.map` / `*-dbg*` files in working tree or in commit 9761a31", "PASS"],
    ["2026-08-13", "GitHub Pages config = main / root, legacy build (via `gh api`)", "PASS"],
    ["2026-08-13", "Live URL reachable: `https://rfaleiros77.github.io/kkml-demo/`", "PASS"],
], [2.4, 10.5, 3.0])

h1("8. Change log")
table(["Date", "Version", "Author", "Changes"], [
    ["2026-08-13", "1.0", "Rogério / Claude", "Initial runbook: environments, four-flavor structure, source-protection rule, local preview, deploy, validation log."],
], [2.4, 1.8, 4.2, 8.1])

doc.add_page_break()
h1("9. Quick Reference — Step Tables")
p("One compact table per procedure. This section is ALWAYS the last one and "
  "is updated together with the prose sections. Prose explains; these tables "
  "execute.")

h2("A. Local preview (details: section 5)")
table(["Step", "Where", "Action"], [
    ["1", "Terminal", "`cd ~/dev/_arquivo/kkml-demo-backup`"],
    ["2", "Terminal", "`python3 -m http.server 8080`"],
    ["3", "Browser", "Open `http://localhost:8080/`"],
    ["4", "Browser", "Click all four cards; confirm each table renders with data"],
    ["5", "Terminal", "Ctrl+C to stop the server"],
], [1.2, 4.6, 10.7])

h2("B. Pre-deploy source-protection check (details: section 4)")
table(["Step", "Where", "Action"], [
    ["1", "Terminal (repo root)", "`find . \\( -name '*.map' -o -name '*-dbg*.js' -o -name '*-dbg*.xml' \\)`"],
    ["2", "Terminal", "Expect NO output. If anything lists, remove it before committing"],
    ["3", "Terminal", "`git status` — confirm no `.map` / `-dbg` file is staged"],
], [1.2, 4.6, 10.7])

h2("C. Deploy to GitHub Pages (details: section 6)")
table(["Step", "Where", "Action"], [
    ["1", "Terminal", "Run check B first (no debug/map files)"],
    ["2", "Terminal", "`git add -A && git commit -m \"<change>\"`"],
    ["3", "Terminal", "`git push origin main`"],
    ["4", "Browser", "Wait ~1–2 min, hard-refresh `https://rfaleiros77.github.io/kkml-demo/`"],
    ["5", "Browser", "Re-test the four flavors from the landing page"],
], [1.2, 4.6, 10.7])

doc.save(OUT)
h1s = [pgh.text for pgh in doc.paragraphs if pgh.style.name.startswith("Heading 1")]
print("written:", OUT)
print("sections:", h1s)
